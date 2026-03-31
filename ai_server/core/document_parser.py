"""
Document Parser for HR AI Server
Извлечение текста из PDF, DOCX, TXT файлов с проверкой безопасности
"""

import os
import io
import re
import base64
import logging
import tempfile
import struct
import zipfile
from typing import Optional, Tuple, List
from pathlib import Path

logger = logging.getLogger(__name__)


class FileSecurityScanner:
    """
    Сканер безопасности файлов.
    Проверяет загружаемые файлы на наличие вредоносного содержимого.
    """

    # Сигнатуры исполняемых файлов (magic bytes)
    EXECUTABLE_SIGNATURES = [
        b'MZ',              # Windows EXE/DLL
        b'\x7fELF',         # Linux ELF
        b'\xfe\xed\xfa',    # macOS Mach-O
        b'\xcf\xfa\xed\xfe', # macOS Mach-O (64-bit)
        b'PK',              # ZIP/JAR (проверяется дополнительно)
    ]

    # Опасные расширения внутри архивов
    DANGEROUS_EXTENSIONS = {
        '.exe', '.bat', '.cmd', '.com', '.scr', '.pif', '.vbs', '.vbe',
        '.js', '.jse', '.wsf', '.wsh', '.ps1', '.psm1', '.msi', '.msp',
        '.dll', '.sys', '.drv', '.cpl', '.inf', '.reg', '.hta', '.jar',
        '.sh', '.bash', '.py', '.rb', '.pl', '.php', '.asp', '.aspx',
    }

    # Паттерны вредоносного содержимого в текстовых/XML файлах
    MALICIOUS_PATTERNS = [
        # JavaScript инъекции
        re.compile(r'<script[\s>]', re.IGNORECASE),
        re.compile(r'javascript:', re.IGNORECASE),
        re.compile(r'on(load|error|click|mouseover)\s*=', re.IGNORECASE),
        # Shell команды
        re.compile(r'(^|\s)(rm\s+-rf|wget\s+|curl\s+.*\|\s*sh|chmod\s+\+x)', re.IGNORECASE),
        # PowerShell
        re.compile(r'powershell\s+.*-enc', re.IGNORECASE),
        re.compile(r'Invoke-(Expression|WebRequest|Command)', re.IGNORECASE),
        # SQL инъекции
        re.compile(r"('\s*(OR|AND)\s+['\d]+\s*=\s*['\d]+)", re.IGNORECASE),
        re.compile(r'(UNION\s+SELECT|DROP\s+TABLE|INSERT\s+INTO|DELETE\s+FROM)', re.IGNORECASE),
        # PHP инъекции
        re.compile(r'<\?php', re.IGNORECASE),
        re.compile(r'(eval|exec|system|passthru|shell_exec)\s*\(', re.IGNORECASE),
        # Обфускация Base64 с исполнением
        re.compile(r'(atob|btoa|base64_decode)\s*\(', re.IGNORECASE),
    ]

    # VBA/Macro паттерны в DOCX
    MACRO_PATTERNS = [
        re.compile(r'Auto(Open|Close|Exec|New)', re.IGNORECASE),
        re.compile(r'(Shell|WScript\.Shell|CreateObject)', re.IGNORECASE),
        re.compile(r'(Environ|Kill|FileCopy|Open.*For\s+Output)', re.IGNORECASE),
    ]

    # Опасные паттерны в PDF
    PDF_DANGEROUS_PATTERNS = [
        re.compile(rb'/JavaScript', re.IGNORECASE),
        re.compile(rb'/JS\s', re.IGNORECASE),
        re.compile(rb'/Launch', re.IGNORECASE),
        re.compile(rb'/OpenAction', re.IGNORECASE),
        re.compile(rb'/AA\s', re.IGNORECASE),      # Additional Actions
        re.compile(rb'/EmbeddedFile', re.IGNORECASE),
        re.compile(rb'/RichMedia', re.IGNORECASE),
        re.compile(rb'/XFA', re.IGNORECASE),         # XML Forms Architecture
    ]

    def scan_file_bytes(self, file_data: bytes, filename: str) -> Tuple[bool, Optional[str]]:
        """
        Комплексная проверка файла на безопасность.

        Returns:
            Tuple[is_safe, threat_description]
        """
        if not file_data:
            return False, "Пустой файл"

        ext = Path(filename).suffix.lower()

        # 1. Проверка magic bytes — файл не является исполняемым?
        is_safe, threat = self._check_magic_bytes(file_data, ext)
        if not is_safe:
            return False, threat

        # 2. Проверка по типу файла
        if ext == '.pdf':
            return self._scan_pdf(file_data)
        elif ext in {'.docx', '.doc'}:
            return self._scan_docx(file_data, ext)
        elif ext in {'.txt', '.rtf'}:
            return self._scan_text(file_data)
        elif ext in {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}:
            return self._scan_image(file_data, ext)

        return True, None

    def _check_magic_bytes(self, data: bytes, expected_ext: str) -> Tuple[bool, Optional[str]]:
        """Проверка что файл соответствует заявленному расширению"""
        # EXE/DLL маскировка
        if data[:2] == b'MZ':
            return False, "Файл является исполняемым (EXE/DLL), загрузка запрещена"
        if data[:4] == b'\x7fELF':
            return False, "Файл является исполняемым (ELF), загрузка запрещена"

        # Проверяем соответствие magic bytes расширению
        if expected_ext == '.pdf' and not data[:5].startswith(b'%PDF'):
            # PDF должен начинаться с %PDF
            if data[:2] == b'PK':
                return False, "Файл является архивом, но имеет расширение PDF"
        elif expected_ext in {'.jpg', '.jpeg'}:
            if data[:2] != b'\xff\xd8':
                return False, "Файл не является валидным JPEG изображением"
        elif expected_ext == '.png':
            if data[:4] != b'\x89PNG':
                return False, "Файл не является валидным PNG изображением"

        return True, None

    def _scan_pdf(self, data: bytes) -> Tuple[bool, Optional[str]]:
        """Сканирование PDF на вредоносное содержимое"""
        threats = []

        for pattern in self.PDF_DANGEROUS_PATTERNS:
            if pattern.search(data):
                threats.append(pattern.pattern.decode('utf-8', errors='ignore').strip('/'))

        if threats:
            logger.warning(f"PDF содержит опасные элементы: {threats}")
            return False, f"PDF содержит потенциально опасные элементы: {', '.join(threats)}"

        return True, None

    def _scan_docx(self, data: bytes, ext: str) -> Tuple[bool, Optional[str]]:
        """Сканирование DOCX на макросы и вредоносное содержимое"""
        # DOCX — это ZIP архив
        if not data[:2] == b'PK':
            if ext == '.doc':
                # Старый формат .doc (OLE2) — проверяем на макросы
                return self._scan_ole_doc(data)
            return False, "Файл повреждён или не является валидным DOCX"

        try:
            with zipfile.ZipFile(io.BytesIO(data)) as zf:
                # Проверяем наличие макросов (vbaProject.bin)
                for name in zf.namelist():
                    name_lower = name.lower()
                    if 'vbaproject' in name_lower or 'macro' in name_lower:
                        return False, "Документ содержит макросы (VBA). Загрузка запрещена"
                    if any(name_lower.endswith(ext) for ext in self.DANGEROUS_EXTENSIONS):
                        return False, f"Документ содержит вложенный опасный файл: {name}"

                # Проверяем XML содержимое на инъекции
                for name in zf.namelist():
                    if name.endswith('.xml') or name.endswith('.rels'):
                        try:
                            content = zf.read(name).decode('utf-8', errors='ignore')
                            for pattern in self.MALICIOUS_PATTERNS:
                                match = pattern.search(content)
                                if match:
                                    logger.warning(f"Обнаружен опасный паттерн в {name}: {match.group()}")
                                    return False, f"Документ содержит подозрительный код в {name}"
                        except Exception:
                            pass

        except zipfile.BadZipFile:
            return False, "Файл повреждён — невалидный DOCX архив"
        except Exception as e:
            logger.error(f"Ошибка сканирования DOCX: {e}")

        return True, None

    def _scan_ole_doc(self, data: bytes) -> Tuple[bool, Optional[str]]:
        """Проверка OLE2 файлов (.doc старого формата)"""
        # OLE2 magic bytes: D0 CF 11 E0 A1 B1 1A E1
        if data[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
            # Ищем VBA строки в бинарном содержимом
            content_str = data.decode('latin-1', errors='ignore')
            for pattern in self.MACRO_PATTERNS:
                if pattern.search(content_str):
                    return False, "Документ .doc содержит макросы (VBA). Загрузка запрещена"
        return True, None

    def _scan_text(self, data: bytes) -> Tuple[bool, Optional[str]]:
        """Сканирование текстовых файлов на инъекции"""
        try:
            text = data.decode('utf-8', errors='ignore')
        except Exception:
            text = data.decode('latin-1', errors='ignore')

        for pattern in self.MALICIOUS_PATTERNS:
            match = pattern.search(text)
            if match:
                logger.warning(f"Обнаружен опасный паттерн в текстовом файле: {match.group()}")
                return False, "Файл содержит подозрительный код или скрипты"

        return True, None

    def _scan_image(self, data: bytes, ext: str) -> Tuple[bool, Optional[str]]:
        """Сканирование изображений на стеганографию и embedded код"""
        # Проверяем что файл реально является изображением
        try:
            from PIL import Image
            img = Image.open(io.BytesIO(data))
            img.verify()
        except ImportError:
            pass  # Pillow не установлен, пропускаем
        except Exception:
            return False, "Файл повреждён или не является валидным изображением"

        # Ищем встроенный исполняемый код после окончания изображения
        if ext in {'.jpg', '.jpeg'}:
            # JPEG заканчивается маркером FFD9
            end_marker = data.rfind(b'\xff\xd9')
            if end_marker > 0 and end_marker < len(data) - 10:
                trailing = data[end_marker + 2:]
                if b'MZ' in trailing or b'<script' in trailing.lower():
                    return False, "Изображение содержит встроенный исполняемый код"

        return True, None


def sanitize_extracted_text(text: str) -> str:
    """
    Очистка извлечённого текста от потенциально опасного содержимого.
    Используется после парсинга документов перед отправкой в AI.
    """
    if not text:
        return text

    # Удаляем HTML теги
    text = re.sub(r'<[^>]+>', '', text)

    # Удаляем JavaScript
    text = re.sub(r'javascript:', '', text, flags=re.IGNORECASE)

    # Удаляем потенциальные инъекции
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.IGNORECASE | re.DOTALL)

    # Удаляем null bytes
    text = text.replace('\x00', '')

    # Удаляем управляющие символы (кроме \n, \r, \t)
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    return text.strip()


class DocumentParser:
    """
    Парсер документов с проверкой безопасности.
    Поддерживает: PDF, DOCX, DOC, TXT, RTF, JPG, PNG, JPEG, BMP, TIFF
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.rtf', '.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}

    def __init__(self, config: dict = None):
        self.config = config or {}
        self.max_file_size = self.config.get('max_file_size_mb', 10) * 1024 * 1024
        self.temp_dir = self.config.get('temp_dir', 'temp')
        self.security_scanner = FileSecurityScanner()

        # Создаём temp директорию
        os.makedirs(self.temp_dir, exist_ok=True)

    def parse_file(self, file_path: str) -> Tuple[str, Optional[str]]:
        """
        Извлечение текста из файла

        Args:
            file_path: Путь к файлу

        Returns:
            Tuple[text, error]: Извлечённый текст и ошибка (если есть)
        """
        path = Path(file_path)

        if not path.exists():
            return "", f"Файл не найден: {file_path}"

        if path.stat().st_size > self.max_file_size:
            return "", f"Файл слишком большой (макс. {self.max_file_size // 1024 // 1024} MB)"

        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return "", f"Неподдерживаемый формат файла: {ext}"

        try:
            if ext == '.pdf':
                return self._parse_pdf(file_path), None
            elif ext in {'.docx', '.doc'}:
                return self._parse_docx(file_path), None
            elif ext == '.txt':
                return self._parse_txt(file_path), None
            elif ext == '.rtf':
                return self._parse_rtf(file_path), None
            elif ext in {'.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif'}:
                return self._parse_image(file_path), None
            else:
                return "", f"Парсер для {ext} не реализован"

        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            return "", str(e)

    def parse_base64(self, content: str, filename: str) -> Tuple[str, Optional[str]]:
        """
        Парсинг файла из base64 с проверкой безопасности

        Args:
            content: Base64 encoded content
            filename: Имя файла (для определения формата)

        Returns:
            Tuple[text, error]
        """
        try:
            # Декодируем base64
            file_data = base64.b64decode(content)

            # --- SECURITY: Проверяем файл перед обработкой ---
            is_safe, threat = self.security_scanner.scan_file_bytes(file_data, filename)
            if not is_safe:
                logger.warning(f"SECURITY: Файл {filename} заблокирован: {threat}")
                return "", f"Файл заблокирован системой безопасности: {threat}"

            # Проверяем размер
            if len(file_data) > self.max_file_size:
                return "", f"Файл слишком большой (макс. {self.max_file_size // 1024 // 1024} MB)"

            # Сохраняем во временный файл
            ext = Path(filename).suffix.lower()
            temp_path = os.path.join(self.temp_dir, f"temp_{os.getpid()}{ext}")

            with open(temp_path, 'wb') as f:
                f.write(file_data)

            try:
                text, error = self.parse_file(temp_path)
                # Санитизация извлечённого текста
                if text:
                    text = sanitize_extracted_text(text)
                return text, error
            finally:
                # Удаляем временный файл
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        except base64.binascii.Error:
            return "", "Невалидный base64 формат"
        except Exception as e:
            logger.error(f"Error parsing base64: {e}")
            return "", str(e)

    def parse_bytes(self, content: bytes, filename: str) -> Tuple[str, Optional[str]]:
        """
        Парсинг файла из bytes с проверкой безопасности

        Args:
            content: Байты файла
            filename: Имя файла

        Returns:
            Tuple[text, error]
        """
        try:
            # --- SECURITY: Проверяем файл перед обработкой ---
            is_safe, threat = self.security_scanner.scan_file_bytes(content, filename)
            if not is_safe:
                logger.warning(f"SECURITY: Файл {filename} заблокирован: {threat}")
                return "", f"Файл заблокирован системой безопасности: {threat}"

            # Проверяем размер
            if len(content) > self.max_file_size:
                return "", f"Файл слишком большой (макс. {self.max_file_size // 1024 // 1024} MB)"

            ext = Path(filename).suffix.lower()
            temp_path = os.path.join(self.temp_dir, f"temp_{os.getpid()}{ext}")

            with open(temp_path, 'wb') as f:
                f.write(content)

            try:
                text, error = self.parse_file(temp_path)
                # Санитизация извлечённого текста
                if text:
                    text = sanitize_extracted_text(text)
                return text, error
            finally:
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        except Exception as e:
            logger.error(f"Error parsing bytes: {e}")
            return "", str(e)

    def _parse_pdf(self, file_path: str) -> str:
        """Извлечение текста из PDF с OCR fallback для сканированных документов"""
        text_parts = []

        # Пробуем pdfplumber (лучше для сложных PDF)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

            if text_parts:
                extracted = '\n\n'.join(text_parts)
                # Проверяем, достаточно ли текста извлечено
                if len(extracted.strip()) >= 100:
                    return extracted
                else:
                    logger.info("PDF содержит мало текста, пробуем OCR...")

        except ImportError:
            logger.warning("pdfplumber not installed, trying PyPDF2")
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")

        # Fallback на PyPDF2
        if not text_parts:
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)

                if text_parts:
                    extracted = '\n\n'.join(text_parts)
                    if len(extracted.strip()) >= 100:
                        return extracted

            except ImportError:
                logger.warning("PyPDF2 not installed")
            except Exception as e:
                logger.warning(f"PyPDF2 failed: {e}")

        # OCR fallback для сканированных PDF
        ocr_text = self._ocr_pdf(file_path)
        if ocr_text:
            return ocr_text

        # Если ничего не помогло, возвращаем что есть
        return '\n\n'.join(text_parts) if text_parts else ""

    def _ocr_pdf(self, file_path: str) -> str:
        """OCR для сканированных PDF"""
        try:
            from pdf2image import convert_from_path
            import pytesseract

            logger.info(f"Запуск OCR для PDF: {file_path}")

            # Конвертируем PDF в изображения
            images = convert_from_path(file_path, dpi=300)

            text_parts = []
            for i, image in enumerate(images):
                # OCR с поддержкой русского и английского
                text = pytesseract.image_to_string(
                    image,
                    lang='rus+eng',
                    config='--psm 1 --oem 3'
                )
                if text.strip():
                    text_parts.append(text)
                logger.debug(f"OCR страница {i+1}: {len(text)} символов")

            if text_parts:
                logger.info(f"OCR успешно: извлечено {sum(len(t) for t in text_parts)} символов")
                return '\n\n'.join(text_parts)

        except ImportError as e:
            logger.warning(f"OCR недоступен (установите pdf2image и pytesseract): {e}")
        except Exception as e:
            logger.error(f"OCR ошибка: {e}")

        return ""

    def _parse_image(self, file_path: str) -> str:
        """OCR для изображений (JPG, PNG, BMP, TIFF)"""
        try:
            import pytesseract
            from PIL import Image

            logger.info(f"Запуск OCR для изображения: {file_path}")

            # Открываем изображение
            image = Image.open(file_path)

            # Предобработка для улучшения OCR
            image = self._preprocess_image(image)

            # OCR с поддержкой русского и английского
            text = pytesseract.image_to_string(
                image,
                lang='rus+eng',
                config='--psm 1 --oem 3'
            )

            if text.strip():
                logger.info(f"OCR успешно: извлечено {len(text)} символов")
                return text.strip()

            return ""

        except ImportError as e:
            logger.error(f"pytesseract не установлен: {e}")
            raise ImportError("Установите pytesseract и Pillow: pip install pytesseract Pillow")
        except Exception as e:
            logger.error(f"Ошибка OCR изображения: {e}")
            raise

    def _preprocess_image(self, image):
        """Предобработка изображения для улучшения OCR"""
        try:
            from PIL import ImageEnhance, ImageFilter

            # Конвертируем в RGB если нужно
            if image.mode != 'RGB':
                image = image.convert('RGB')

            # Увеличиваем контраст
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)

            # Увеличиваем резкость
            image = image.filter(ImageFilter.SHARPEN)

            # Масштабируем если слишком маленькое
            width, height = image.size
            if width < 1000:
                ratio = 1000 / width
                image = image.resize((int(width * ratio), int(height * ratio)))

            return image

        except Exception as e:
            logger.warning(f"Предобработка изображения не удалась: {e}")
            return image

    def _parse_docx(self, file_path: str) -> str:
        """Извлечение текста из DOCX"""
        try:
            from docx import Document
            doc = Document(file_path)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            return '\n'.join(text_parts)

        except ImportError:
            raise ImportError("Установите python-docx: pip install python-docx")

    def _parse_txt(self, file_path: str) -> str:
        """Чтение текстового файла"""
        encodings = ['utf-8', 'cp1251', 'cp1252', 'latin-1']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue

        # Если ничего не помогло, читаем с игнорированием ошибок
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _parse_rtf(self, file_path: str) -> str:
        """Извлечение текста из RTF"""
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            return rtf_to_text(rtf_content)

        except ImportError:
            # Простой fallback - убираем RTF теги
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()

            import re
            # Убираем RTF команды
            content = re.sub(r'\\[a-z]+\d*\s?', '', content)
            content = re.sub(r'[{}]', '', content)
            return content.strip()

    def get_supported_formats(self) -> list:
        """Список поддерживаемых форматов"""
        return list(self.SUPPORTED_EXTENSIONS)

    def validate_file(self, filename: str, file_size: int) -> Tuple[bool, Optional[str]]:
        """
        Валидация файла перед парсингом

        Args:
            filename: Имя файла
            file_size: Размер в байтах

        Returns:
            Tuple[is_valid, error_message]
        """
        # Защита от path traversal
        basename = Path(filename).name
        if '..' in filename or '/' in filename or '\\' in filename:
            if basename != filename:
                logger.warning(f"SECURITY: Попытка path traversal: {filename}")

        ext = Path(basename).suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Неподдерживаемый формат: {ext}. Поддерживаются: {', '.join(self.SUPPORTED_EXTENSIONS)}"

        # Двойные расширения (например, file.pdf.exe)
        name_parts = basename.split('.')
        if len(name_parts) > 2:
            for part in name_parts[1:]:
                if f'.{part.lower()}' in FileSecurityScanner.DANGEROUS_EXTENSIONS:
                    return False, f"Файл имеет опасное расширение: .{part}"

        if file_size > self.max_file_size:
            max_mb = self.max_file_size // 1024 // 1024
            return False, f"Файл слишком большой. Максимум: {max_mb} MB"

        if file_size == 0:
            return False, "Файл пустой"

        return True, None
